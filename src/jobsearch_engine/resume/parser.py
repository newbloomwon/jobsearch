"""Resume file → plain text.

PDF/DOCX libraries are optional dependencies (the `resume` extra), imported
lazily so the engine core has no heavy imports when they're not installed.
"""
from __future__ import annotations

import io

MAX_RESUME_BYTES = 5 * 1024 * 1024


class ResumeParseError(ValueError):
    pass


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded resume (pdf/docx/txt/md)."""
    if not data:
        raise ResumeParseError("Empty file")
    if len(data) > MAX_RESUME_BYTES:
        raise ResumeParseError("File exceeds the 5MB limit")

    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith((".txt", ".md", ".text", ".rtf")) or not name:
        return _from_plain(data)
    # Unknown extension: sniff docx (zip magic) then fall back to plain text.
    if data[:2] == b"PK":
        return _from_docx(data)
    if data[:4] == b"%PDF":
        return _from_pdf(data)
    return _from_plain(data)


def _from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - env dependent
        raise ResumeParseError(
            "PDF parsing requires the 'resume' extra: pip install '.[resume]'"
        ) from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ResumeParseError(f"Could not read PDF: {exc}") from exc
    text = "\n".join(pages).strip()
    if not text:
        raise ResumeParseError(
            "PDF contained no extractable text (it may be a scanned image)"
        )
    return text


def _from_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - env dependent
        raise ResumeParseError(
            "DOCX parsing requires the 'resume' extra: pip install '.[resume]'"
        ) from exc
    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
    except Exception as exc:
        raise ResumeParseError(f"Could not read DOCX: {exc}") from exc
    return "\n".join(parts).strip()


def _from_plain(data: bytes) -> str:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
    text = text.strip()
    if not text:
        raise ResumeParseError("File contained no text")
    return text
